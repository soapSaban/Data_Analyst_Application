import os
import streamlit as st
import pandas as pd
import pdfplumber
from docx import Document as DocxDocument
from PIL import Image
import pytesseract
import matplotlib.pyplot as plt
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from dotenv import load_dotenv
import requests
from textwrap import wrap

# LangChain and Embeddings
from langchain_community.vectorstores import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.docstore.document import Document

# Configure Tesseract path (UPDATE FOR YOUR SYSTEM)
pytesseract.pytesseract.tesseract_cmd = r'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'  # Windows
# pytesseract.pytesseract.tesseract_cmd = '/usr/local/bin/tesseract'  # Mac/Linux

# Load environment variables
load_dotenv()
TOGETHER_API_KEY = os.getenv("TOGETHER_API_KEY")

# Initialize embeddings safely
try:
    embedding_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
except Exception as e:
    st.error(f"Failed to initialize embeddings: {str(e)}")
    st.stop()

# Streamlit config
st.set_page_config(page_title="AI Data Analyst", layout="wide")
st.title("📊 AI Data Analyst by Tanmay")
st.write("Upload files and interact with your data using an AI agent powered by Together.ai.")

# Check for API key
if not TOGETHER_API_KEY:
    st.error("⚠️ TOGETHER_API_KEY not found! Please add it to your .env file.")
    st.stop()

# Tabs
tab1, tab2, tab3, tab4 = st.tabs(["📄 Summary", "📊 Visualizations", "💬 Chat", "⬇ Download Report"])

# Session state
if 'summary' not in st.session_state:
    st.session_state.summary = ""
if 'insights' not in st.session_state:
    st.session_state.insights = ""
if 'dataframes' not in st.session_state:
    st.session_state.dataframes = []
if 'vectordb' not in st.session_state:
    st.session_state.vectordb = None
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []

# --- Extract text from files ---
def extract_text(file):
    text = ""
    try:
        if file.name.endswith(".pdf"):
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    else:
                        try:
                            img = page.to_image(resolution=300).original.convert("L")
                            text += pytesseract.image_to_string(img) + "\n"
                        except Exception as img_e:
                            st.warning(f"Couldn't extract text from PDF image: {str(img_e)}")
        elif file.name.endswith(".docx"):
            doc = DocxDocument(file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif file.name.endswith(".txt"):
            try:
                text += file.read().decode("utf-8")
            except UnicodeDecodeError:
                text += file.read().decode("latin-1")
        elif file.name.endswith((".csv", ".xlsx")):
            try:
                df = pd.read_csv(file) if file.name.endswith(".csv") else pd.read_excel(file)
                if not any(name == file.name for name, _ in st.session_state.dataframes):
                    st.session_state.dataframes.append((file.name, df))
                text += df.to_string()
            except Exception as df_e:
                st.error(f"Error reading data file {file.name}: {str(df_e)}")
        elif file.name.lower().endswith((".jpg", ".jpeg", ".png")):
            try:
                image = Image.open(file).convert("L")
                st.image(image, caption="Uploaded Image", width=300)  # Show preview
                text += pytesseract.image_to_string(image)
                st.success("Text extracted successfully!")
            except Exception as img_e:
                st.error(f"⚠️ OCR Error: {str(img_e)}. Please ensure Tesseract is installed.")
    except Exception as e:
        st.error(f"Error reading {file.name}: {str(e)}")
    return text

# --- Generate charts ---
def generate_chart(df):
    charts = []
    numeric_cols = df.select_dtypes(include="number").columns
    if len(numeric_cols) == 0:
        st.warning("No numeric columns found for visualization.")
        return charts
    
    for col in numeric_cols[:2]:  # Limit to first 2 numeric columns
        try:
            fig, ax = plt.subplots(figsize=(10, 6))
            df[col].value_counts().head(10).plot(kind="bar", ax=ax, title=f"Top 10 values of {col}")
            plt.xticks(rotation=45)
            plt.tight_layout()
            buf = BytesIO()
            plt.savefig(buf, format="png", dpi=100)
            buf.seek(0)
            charts.append(buf)
            plt.close(fig)
        except Exception as e:
            st.error(f"Error generating chart for column {col}: {str(e)}")
    return charts

# --- Together.ai API call ---
def together_ai(prompt, max_tokens=500):
    url = "https://api.together.xyz/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {TOGETHER_API_KEY}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": max_tokens,
        "temperature": 0.7,
        "top_p": 0.9,
        "stream": False
    }
    
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=30)
        
        if response.status_code != 200:
            st.error(f"API Error {response.status_code}: {response.text}")
            return f"API Error: {response.status_code}"
        
        result = response.json()
        
        if "choices" in result and len(result["choices"]) > 0:
            content = result["choices"][0].get("message", {}).get("content", "")
            if content:
                return content.strip()
            return "No content in API response."
        else:
            st.error(f"Unexpected API response format: {result}")
            return "API returned unexpected format."
            
    except requests.exceptions.RequestException as e:
        st.error(f"Network Error: {str(e)}")
        return "Network error occurred."
    except Exception as e:
        st.error(f"API Error: {str(e)}")
        return "Error generating response."

# --- File upload ---
uploaded_files = st.file_uploader(
    "Upload your files",
    type=["pdf", "docx", "txt", "csv", "xlsx", "jpg", "png", "jpeg"],
    accept_multiple_files=True
)

if uploaded_files:
    all_text = ""
    for file in uploaded_files:
        file_text = extract_text(file)
        if file_text:
            all_text += file_text + "\n\n"
    
    if all_text.strip():
        try:
            splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
            chunks = splitter.split_text(all_text)
            docs = [Document(page_content=chunk) for chunk in chunks]

            st.session_state.vectordb = Chroma.from_documents(
                documents=docs,
                embedding=embedding_model,
                persist_directory=None
            )
        except Exception as e:
            st.error(f"Error creating vector database: {str(e)}")

    # --- Summary Tab ---
    with tab1:
        st.subheader("📄 Document Summary")
        if not st.session_state.summary and all_text.strip():
            with st.spinner("Generating summary..."):
                summary_prompt = f"""Please provide a clear and concise summary of the following content.
                Focus on key points and main ideas. Keep it under 300 words.
                
                Content:
                {all_text[:10000]}"""
                st.session_state.summary = together_ai(summary_prompt)
        st.write(st.session_state.summary)

        if not st.session_state.insights and all_text.strip():
            with st.spinner("Generating insights..."):
                insights_prompt = f"""Based on the following content, provide 3-5 key insights and 
                actionable recommendations. Be specific and practical.
                
                Content:
                {all_text[:10000]}"""
                st.session_state.insights = together_ai(insights_prompt)
        st.write("### 🔍 Key Insights & Recommendations")
        st.write(st.session_state.insights)

# --- Visualization Tab ---
with tab2:
    st.subheader("📊 Auto-Generated Visualizations")
    if st.session_state.dataframes:
        for name, df in st.session_state.dataframes:
            st.write(f"### {name}")
            st.dataframe(df.head())
            
            charts = generate_chart(df)
            if charts:
                cols = st.columns(2)
                for i, chart in enumerate(charts[:2]):  # Show max 2 charts
                    with cols[i % 2]:
                        st.image(chart)
                        st.download_button(
                            f"Download {name} Chart {i+1}",
                            chart.getvalue(),
                            file_name=f"{name.split('.')[0]}_chart_{i+1}.png",
                            mime="image/png"
                        )
    else:
        st.info("Upload CSV or Excel files to see visualizations")

# --- Chat Tab ---
with tab3:
    st.subheader("💬 Ask Questions About Your Data")
    question = st.text_input("Enter your question:", key="question_input")
    
    if st.button("Ask") and question:
        if st.session_state.vectordb:
            try:
                docs = st.session_state.vectordb.similarity_search(question, k=3)
                context = "\n".join([d.page_content for d in docs])
                chat_prompt = f"""Based on this context:
                {context}
                
                Please answer this question: {question}
                - Be concise but thorough
                - If unsure, say you don't know
                - Format your answer clearly"""
                
                with st.spinner("Generating answer..."):
                    answer = together_ai(chat_prompt)
                    st.session_state.chat_history.append((question, answer))
            except Exception as e:
                st.error(f"Error searching documents: {str(e)}")
                answer = "Sorry, there was an error processing your question."
        else:
            st.warning("Please upload and process files first.")
            answer = None
        
        if answer:
            st.write(f"**Answer:** {answer}")

    if st.session_state.chat_history:
        with st.expander("Chat History"):
            for i, (q, a) in enumerate(st.session_state.chat_history):
                st.markdown(f"**Q{i+1}:** {q}")
                st.markdown(f"**A{i+1}:** {a}")
                st.divider()

# --- Report Download Tab ---
with tab4:
    st.subheader("⬇ Download AI Report")
    if st.button("Generate PDF Report"):
        if not st.session_state.summary and not st.session_state.insights:
            st.warning("Please generate summary and insights first by uploading files.")
        else:
            try:
                buffer = BytesIO()
                c = canvas.Canvas(buffer, pagesize=letter)
                width, height = letter

                # Title
                c.setFont("Helvetica-Bold", 16)
                c.drawString(72, height - 72, "AI Data Analyst Report")
                c.line(72, height - 80, width - 72, height - 80)

                # Summary section
                c.setFont("Helvetica-Bold", 14)
                c.drawString(72, height - 110, "Summary:")
                c.setFont("Helvetica", 12)
                text = c.beginText(72, height - 130)
                for line in wrap(st.session_state.summary, width=100)[:40]:  # Limit lines
                    text.textLine(line)
                c.drawText(text)

                # Insights section
                c.setFont("Helvetica-Bold", 14)
                c.drawString(72, height - 490, "Key Insights:")
                c.setFont("Helvetica", 12)
                text = c.beginText(72, height - 510)
                for line in wrap(st.session_state.insights, width=100)[:40]:  # Limit lines
                    text.textLine(line)
                c.drawText(text)

                # Add page for data preview if available
                if st.session_state.dataframes:
                    c.showPage()
                    c.setFont("Helvetica-Bold", 14)
                    c.drawString(72, height - 72, "Data Preview")
                    
                    y_position = height - 100
                    for name, df in st.session_state.dataframes[:2]:  # Limit to 2 dataframes
                        c.setFont("Helvetica-Bold", 12)
                        c.drawString(72, y_position, f"Data from: {name}")
                        y_position -= 20
                        
                        c.setFont("Helvetica", 10)
                        text = c.beginText(72, y_position)
                        for line in str(df.head()).split('\n')[:10]:  # Limit lines
                            text.textLine(line)
                        c.drawText(text)
                        y_position -= 150
                        if y_position < 100:
                            c.showPage()
                            y_position = height - 72

                c.save()
                buffer.seek(0)
                
                st.download_button(
                    "Download Full Report",
                    buffer,
                    file_name="AI_Data_Analyst_Report.pdf",
                    mime="application/pdf"
                )
            except Exception as e:
                st.error(f"Error generating PDF: {str(e)}")

# Add some spacing at the bottom
st.markdown("<br><br>", unsafe_allow_html=True)

